import time
import shutil
import os
import sys
import config
import path
import webbrowser
import docx
from jira import JIRA
from datetime import datetime
from genericpath import isdir


def logging_deco(func):
    def wrapped_func(*args):
        start_r = time.perf_counter()
        start_p = time.process_time()
        ret = func(*args)
        end_r = time.perf_counter()
        end_p = time.process_time()
        
        elapsed_r = end_r - start_r
        elapsed_p = end_p - start_p
        print(f'{func.__name__} : {elapsed_r:.6f}sec (Perf_Counter) / {elapsed_p:.6f}sec (Process Time)')
        return ret
    return wrapped_func

def Exception_decorator(func):
    def wrapped_func(*args):
        try:
            ret = func(*args)
        except Exception as e:
            print(e)
        return ret
    return wrapped_func

class Jira:
    def __init__(self):
        # Jira Connection
        self.url_prefix = config.jira_server + "/browse/"
        self.options = {'server': config.jira_server}
        self.jira = JIRA(self.options, basic_auth=config.jira_basic_auth)
        self.myid = config.jira_myid
        self.today_yyyymmdd = datetime.now().strftime('%Y-%m-%d')

    def get_deployInfo(self, target_path:str, site_code:str) -> tuple[str, str, list]:
        """ 
        target_path 내부 파일들의 정보를 가져옴.
        make_template 안에서 사용됨.

        Args:
        - target_path : 
        - site_code : 사이트
        
        Returns:
        - siteVersion : 
        - changeName : 
        - filePathList : 
        """
        site_version = ""
        change_name = ""
        seperator_string = "-r" if "." in site_code else ".r"
        target_path += "/" + site_code
        list_filepath = [] 

        for x in os.listdir(target_path):
            if ".rpm" in x:
                site_version += x.split(seperator_string)[0]+"\n"
                list_filepath.append(target_path + "/" + x)

            if ".txt" in x:
                change_name = x

        return site_version, change_name, list_filepath

    def make_template(self, site_code:str, workspace:str) -> tuple[str, list]:
        """ 
        Comment : 지라 댓글 탬플릿 함수
        Author : 
        Args:
        - site_code : 사이트 코드
        - workspace : 작업 경로, path.file_path를 기본으로 사용
        
        Returns:
        - jira_template : 지라 댓글에 남길 문자열
        - jira_issue_link : txt에서 검출된 지라 이슈 코드들
        """
        target = path.dict_company[site_code]["name"]
        download_link = path.dict_company[site_code]["url"]
        
        site_version, change_name, _ = self.get_deployInfo(workspace, site_code)

        site_contents = ""
        txt_filename = f"{workspace}/{site_code}/{change_name}"

        try:
            with open(txt_filename, "r", encoding="utf8") as file:
                for tmp in file.readlines():
                    site_contents += tmp.replace("#", "")
        except Exception as e:
            print(e) # 에러처리용 데코레이터 추가
            return "", []



        jira_issue_link = []
        for x in [tmp.split("]")[0] for tmp in site_contents.replace(" ", "").split("[") if "]" in tmp]:
            print(x)
            if "-" not in x:
                continue
            jira_issue_link.append(x)

        jira_template = "*" + target 
        jira_template += f"({site_code}) " if "." not in site_code else " "
        jira_template += "RPM 배포*\n"
        jira_template += f"\n버전 : \n{site_version} \n"
        jira_template += f"\n수정 내용 : \n{site_contents}\n"
        jira_template += f"\n다운로드 경로 : \n{download_link}\n"

        return jira_template, jira_issue_link

    def add_comment(self, issue_code:str, comment:str) -> None: 
        """ 
        issue_code에 댓글 추가 

        Args:
        - issue_code : 이슈 코드
        - comment : 댓글
        """
        comment_to_edit = self.jira.add_comment(issue_code, 'Change this content later')
        comment_to_edit.update(body=comment)

    def append_docx(self, file_name:str, paragraph:str) -> None:
        """ 구글 드라이브 패치 내용.docx에 txt 파일 내용 추가 """
        try:
            doc = docx.Document(file_name)
        except:
            doc = docx.Document()
            doc.save(file_name)
            doc = docx.Document(file_name)

        doc.add_paragraph().add_run("\n" + datetime.now().strftime('%Y-%m-%d')).bold = True
        doc.add_paragraph(paragraph)
        doc.save(file_name)
        
    @logging_deco
    def upload_gdrive(self, site_code:str, workspace:str) -> None:
        """ site_code(폴더명)에 .이 있으면 버전 배포로 인식 """
        target_path=""
        if "." not in site_code:
            for x in os.listdir(path.gdirve_path):
                if site_code + "_" in x:
                    target_path = path.gdirve_path + "/" + x + "/패치"
                    break
        else:
            for x in os.listdir(path.gdrive_path_version):
                mid_path = ".".join(map(str, site_code.split(".")[:-1]))
                if mid_path in x:
                    target_path = path.gdrive_path_version + f"/V{mid_path}/{site_code}/패치"
                    break

        print(target_path)
        _, change_name, list_filepath = self.get_deployInfo(workspace, site_code)

        for x in list_filepath:
            try:
                shutil.copy(x, target_path)
                print(x + " - 업로드 완료")
            except:
                print(x + " - 실패")

        webbrowser.open(path.dict_company[site_code]["url"])

        site_contents = ""
        try:
            with open(workspace + "/" + site_code + "/" + change_name, "r", encoding="utf8") as file:
                for tmp in file.readlines():
                    site_contents += tmp.replace("#", "")

            print("파일 명 : " + target_path + path.gdrive_patch_docx)
            self.append_docx(target_path + path.gdrive_patch_docx, site_contents + "\n")
        except:
            pass

    @logging_deco
    def auto_comment(self, site_code:str, white_list:list, workspace:str) -> str:
        """ 지라 이슈에 댓글 달기 """
        if path.dict_company[site_code]["url"] == "":
            return "다운로드 링크 없음"

        if path.dict_company[site_code]["name"] == "":
            return "사이트명 없음"

        if site_code != "":
            str_template, issue_code_list = self.make_template(site_code, workspace)
        else:
            return "사이트 코드 없음"
        
        if issue_code_list == []:
            return "No Issue"

        for issue_code in list(set(issue_code_list)):
            if issue_code in white_list or "ALL" in white_list:
                self.add_comment(issue_code, str_template)
                print(config.jira_server + "/browse/" + issue_code)
                webbrowser.open(config.jira_server + "/browse/" + issue_code)

        return "성공"


if __name__ == '__main__':
    """
    RPM 배포 자동화 스크립트
    - 최초 작성일 : 2022-09-15
    - 마지막 수정일 : 2022-11-29
    - 마지막 사용일 : 2022-11-29 성공
    - 사용 조건 : 
        1. workspace 위치에 컴퍼니 코드(버전)을 이름으로 가진 폴더(수정내용.txt, .rpm 파일) 존재
        2. path.py에 해당 컴퍼니 코드(버전)의 정보 존재
        3. config.py에 id와 token값이 기입되어 있고, Jira 시스템 상에서도 유효해야 함
    - 주의 사항 : 
        1. white_list 에 값이 "ALL" 이면 txt에서 검출된 모든 이슈코드에 댓글이 올라감. 
        2. txt파일이 있고, 수정 사항 중 일부 이슈에만 댓글을 남길 경우 white_list에 이슈코드 직접 기입.
            ex) white_list = ["H0000-0000", "H1234-1234"]
        3. txt 파일이 없을 경우 지라 댓글과 패치 내용(구글 드라이브)이 남지 않음
        4. path.py에 컴퍼니 코드가 없을 경우 작업에서 생략됨
        5. 작업이 완료된 대상은 path.after_upload_path + "/" + jira.today_yyyymmdd 경로로 이동됨
    """
    white_list = ["ALL"]
    
    workspace = path.file_path
    if not isdir(workspace):
        print(f"Error : {workspace} - Directory Not Found")
        sys.exit()

    site_codes = []
    for x in os.listdir(workspace):
        if isdir(workspace + "/" + x) and x in [key for key in path.dict_company]:
            site_codes.append(x)

    jira = Jira()
    if not isdir(workspace):
        os.mkdir(path.after_upload_path + "/" + jira.today_yyyymmdd)

    for site_code in site_codes:
        jira.upload_gdrive(site_code, workspace)
        print(jira.auto_comment(site_code, white_list, workspace))
        shutil.move(workspace + "/" + site_code, path.after_upload_path + "/" + jira.today_yyyymmdd + "/" + site_code)
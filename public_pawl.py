import time
import shutil
import os
import sys
from datetime import datetime
import time
from genericpath import isdir

import webbrowser
import docx
import pyexcel
from jira import JIRA
import paramiko

from distutils.command.build import build
import json
from urllib.request import HTTPBasicAuthHandler, HTTPPasswordMgrWithDefaultRealm, build_opener, install_opener, urlopen
import requests
import winsound as ws
import ctypes

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.alert import Alert
from selenium.webdriver.common.by import By

import pyautogui
import pywinauto

import path
import config


def TEST(func):
    def wrapped_func(*args):
        print("*********************테스트 중인 함수입니다.*********************")
        ret = func(*args)
        print("*********************결과를 보장할 수 없음.*********************")
        return ret
    return wrapped_func

def logging_deco(func):
    def wrapped_func(*args):
        start_r = time.perf_counter()
        start_p = time.process_time()
        ret = func(*args)
        end_r = time.perf_counter()
        end_p = time.process_time()
        
        elapsed_r = end_r - start_r
        elapsed_p = end_p - start_p
        print(f'{func.__name__} : {elapsed_r:.6f}sec (Perf_Counter) / {elapsed_p:.6f}sec (Process Time)')
        return ret
    return wrapped_func

def Exception_decorator(func):
    def wrapped_func(*args):
        try:
            ret = func(*args)
        except Exception as e:
            print(e)
        return ret
    return wrapped_func


class Jira:
    def __init__(self):
        # Jira Connection
        self.url_prefix = config.jira_server + "/browse/"
        self.options = {'server': config.jira_server}
        self.jira = JIRA(self.options, basic_auth=config.jira_basic_auth)
        self.myid = config.jira_myid
        self.today_yyyymmdd = datetime.now().strftime('%Y-%m-%d')

    def get_deployInfo(self, target_path:str, site_code:str) -> tuple[str, str, list]:
        """ 
        target_path 내부 파일들의 정보를 가져옴.
        make_template 안에서 사용됨.

        Args:
        - target_path : 
        - site_code : 사이트
        
        Returns:
        - siteVersion : 
        - changeName : 
        - filePathList : 
        """
        site_version = ""
        change_name = ""
        seperator_string = "-r" if "." in site_code else ".r"
        target_path += "/" + site_code
        list_filepath = [] 

        for x in os.listdir(target_path):
            if ".rpm" in x:
                site_version += x.split(seperator_string)[0]+"\n"
                list_filepath.append(target_path + "/" + x)

            if ".txt" in x:
                change_name = x

        return site_version, change_name, list_filepath

    def make_template(self, site_code:str, workspace:str) -> tuple[str, list]:
        """ 
        Comment : 지라 댓글 탬플릿 함수
        Author : 
        Args:
        - site_code : 사이트 코드
        - workspace : 작업 경로, path.file_path를 기본으로 사용
        
        Returns:
        - jira_template : 지라 댓글에 남길 문자열
        - jira_issue_link : txt에서 검출된 지라 이슈 코드들
        """
        target = path.dict_company[site_code]["name"]
        download_link = path.dict_company[site_code]["url"]
        
        site_version, change_name, _ = self.get_deployInfo(workspace, site_code)

        site_contents = ""
        txt_filename = f"{workspace}/{site_code}/{change_name}"

        try:
            with open(txt_filename, "r", encoding="utf8") as file:
                for tmp in file.readlines():
                    site_contents += tmp.replace("#", "")
        except Exception as e:
            print(e) # 에러처리용 데코레이터 추가
            return "", []



        jira_issue_link = []
        for x in [tmp.split("]")[0] for tmp in site_contents.replace(" ", "").split("[") if "]" in tmp]:
            print(x)
            if "-" not in x:
                continue
            jira_issue_link.append(x)

        jira_template = "*" + target 
        jira_template += f"({site_code}) " if "." not in site_code else " "
        jira_template += "RPM 배포*\n"
        jira_template += f"\n버전 : \n{site_version} \n"
        jira_template += f"\n수정 내용 : \n{site_contents}\n"
        jira_template += f"\n다운로드 경로 : \n{download_link}\n"

        return jira_template, jira_issue_link

    def add_comment(self, issue_code:str, comment:str) -> None: 
        """ 
        issue_code에 댓글 추가 

        Args:
        - issue_code : 이슈 코드
        - comment : 댓글
        """
        comment_to_edit = self.jira.add_comment(issue_code, 'Change this content later')
        comment_to_edit.update(body=comment)

    def append_docx(self, file_name:str, paragraph:str) -> None:
        """ 구글 드라이브 패치 내용.docx에 txt 파일 내용 추가 """
        try:
            doc = docx.Document(file_name)
        except:
            doc = docx.Document()
            doc.save(file_name)
            doc = docx.Document(file_name)

        doc.add_paragraph().add_run("\n" + datetime.now().strftime('%Y-%m-%d')).bold = True
        doc.add_paragraph(paragraph)
        doc.save(file_name)
        
    @logging_deco
    def upload_gdrive(self, site_code:str, workspace:str) -> None:
        """ site_code(폴더명)에 .이 있으면 버전 배포로 인식 """
        target_path=""
        if "." not in site_code:
            for x in os.listdir(path.gdirve_path):
                if site_code + "_" in x:
                    target_path = path.gdirve_path + "/" + x + "/패치"
                    break
        else:
            for x in os.listdir(path.gdrive_path_version):
                mid_path = ".".join(map(str, site_code.split(".")[:-1]))
                if mid_path in x:
                    target_path = path.gdrive_path_version + f"/V{mid_path}/{site_code}/패치"
                    break

        print(target_path)
        _, change_name, list_filepath = self.get_deployInfo(workspace, site_code)

        for x in list_filepath:
            try:
                shutil.copy(x, target_path)
                print(x + " - 업로드 완료")
            except:
                print(x + " - 실패")

        webbrowser.open(path.dict_company[site_code]["url"])

        site_contents = ""
        try:
            with open(workspace + "/" + site_code + "/" + change_name, "r", encoding="utf8") as file:
                for tmp in file.readlines():
                    site_contents += tmp.replace("#", "")

            print("파일 명 : " + target_path + path.gdrive_patch_docx)
            self.append_docx(target_path + path.gdrive_patch_docx, site_contents + "\n")
        except:
            pass

    @logging_deco
    def auto_comment(self, site_code:str, white_list:list, workspace:str) -> str:
        """ 지라 이슈에 댓글 달기 """
        if path.dict_company[site_code]["url"] == "":
            return "다운로드 링크 없음"

        if path.dict_company[site_code]["name"] == "":
            return "사이트명 없음"

        if site_code != "":
            str_template, issue_code_list = self.make_template(site_code, workspace)
        else:
            return "사이트 코드 없음"
        
        if issue_code_list == []:
            return "No Issue"

        for issue_code in list(set(issue_code_list)):
            if issue_code in white_list or "ALL" in white_list:
                self.add_comment(issue_code, str_template)
                print(config.jira_server + "/browse/" + issue_code)
                webbrowser.open(config.jira_server + "/browse/" + issue_code)

        return "성공"

    @TEST
    def _search_issue(self, project, repoter): # jql을 사용한 이슈 검색
        search_issue_jql = f"summary~AsyncError and project={project} and search_issue={repoter} and status not in (closed, done)"
        jira_issue = self.jira.search_issues(search_issue_jql)
        return jira_issue

    @TEST
    def get_comments(self, issue_code): # 특정 이슈 댓글 가져오기
        return self.comments(issue_code)

    @TEST
    def get_assignee(self, task_key): # 해당 이슈의 담당자를 얻어옴
        issue = self.jira.issue(task_key)
        return issue.fields.assignee.name

    @TEST
    def set_asignee(self): # 특정 이슈의 담당자 변경
        pass

    @TEST 
    def get_repoter(self): # 특정 이슈의 보고자 얻어오기
        pass

    @TEST 
    def set_issue_statue(self): # 이슈 상태 변경
        pass

    @TEST
    def get_issue_status(self): # 이슈 상태 얻기
        pass
    

    def get_projects(self): # 프로젝트 전체 목록 얻어오기 2022-09-08 테스트 성공
        return self.jira.projects()

    @logging_deco
    def get_my_issue(self) -> list:
        """ 
        나(config.py의 myid)에게 할당된 이슈의 정보를 정리하여 반환함.

        Returns:
        - list_res : [[지난 날짜, 이슈 코드, 패치 일정, 고객사, 내용 요약, url], ...]
        """
        list_res = []
        for x in self.list_projects_key:
            q = f'project = "{x}" AND assignee = {self.myid} ORDER BY created DESC'
            try:
                for iss in self.jira.search_issues(q):
                    day = datetime.now() - datetime.strptime(iss.fields.created, '%Y-%m-%dT%H:%M:%S.%f%z').replace(tzinfo=None)
                    list_tmp = [day.days, iss.key, iss.fields.customfield_10172, iss.fields.customfield_10121, iss.fields.summary, config.jira_server + "/browse/" + iss.key]
                    
                    list_res.append(list_tmp)
            except:
                pass

        return list_res

class Ssh():
    def __init__(self):
        pass

    # 컴퍼니 브랜치 repo 시 사용
    def ssh_cmd(ip, port, user, pw, cmds):
        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy)
        ssh.connect(ip, port, user, pw)

        stdin, stdout, stderr = ssh.exec_command(";".join(cmds))

        lines = stdout.read()
        res = ''.join(str(lines))

        return res

    def ssh_cmd_backup(ip, port, user, pw, cmds):
        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy)
        ssh.connect(ip, port, user, pw)

        stdin, stdout, stderr = ssh.exec_command(";".join(cmds))
        #ch = ssh.get_transport().open_session()
        #res = ch.recv_exit_status()
        #stdin.write(cmd+'\n')
        #stdin.flush()
        #stdout._set_mode('b')

        lines = stdout.read()
        res = ''.join(str(lines))

        '''
        for line in lines:
            re = str(line).replace('\n', '')
            print(re)
            if str(re) in "fin":
                print("fin")
                break;
        '''

        return res


    def wait_stems(ch):
        time.sleep(1)
        outdata = errdata = ""

        while ch.recv_ready():
            outdata += str(ch.recv(1024))

        while ch.recv_stderr_ready():
            errdata += str(ch.recv_stderr(1024))

        return outdata, errdata


    def ssh_connection(ip, port, user, pw):
        try:
            ssh = paramiko.SSHClient()
            ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy)
            ssh.connect(ip, port, user, pw)

            console = ssh.invoke_shell()
            console.settimeout(9999)
            console.keep_this = ssh
            time.sleep(3)

            return console

        except Exception as err:
            print(err)


    # 원격으로 관리 대상 서버의 프로퍼티를 체크하여 이상 유무를 파악
    def ssh_property_check(cmd):
        servers = config.ssh_servers

        for server in servers:
            print(server["ip"], server["port"], server["id"], server["pw"], cmd)

class Selenium():
    def __init__(self):
        chrome_option = Options()
        chrome_option.add_argument('--headless')
        chrome_option.add_argument('--log-level=3')
        chrome_option.add_argument('--disable-logging')
        chrome_option.add_argument('--no-sandbox')
        chrome_option.add_argument('--disable-gpu')

        chrome_option.add_argument('--ignore-certificate-errors')
        chrome_option.add_argument('--ignore-ssl-errors')

        chrome_driver = webdriver.Chrome(executable_path='./chromedriver.exe', chrome_options = chrome_option)
        self.chrome_driver = chrome_driver

        #self.company_branch_url = f"{config.gerrit_host}/gitweb?p=onenet/company_code.git;a=blob;f=ftc_companyName.csv;h=802c5cc6b2867547b4486d21fd8ec1d2e98e327b;hb=refs/heads/master"

    def selenium_1(self):
        self.chrome_driver.get(self.company_branch_url)

        self.chrome_driver.quit()


    def check_branchName(self, str_company):
        self.chrome_driver.get(self.company_branch_url)

        time.sleep(3000) 
        self.chrome_driver.switch_to.window()

        try:
            input_id = self.chrome_driver.find_element_by_xpath('//*[@id="j_username"]')
            input_id.send_keys('test')

        except Exception as err:
            print(err)

        btn_login = self.chrome_driver.find_element_by_css_selector('input.btn.btn-primary')
        btn_login.click()

        time.sleep(3000)

        return False


    def test_check_alert(self, id, pw):
        baseurl = config.gerrit_host + "/login/#/"

        wait = WebDriverWait(self.chrome_driver, 30)

        self.chrome_driver.get(baseurl)

        time.sleep(3)

        #webdriver.ActionChains(dr).send_keys(id+Keys.TAB).perform()

        #pyautogui.moveTo(1000, 700, 1)

        #alert = wait.until(EC.alert_is_present())

        #alert.send_keys(id)
        #alert.send_keys(Keys.TAB)
        #alert.send_keys(pw)

        #alert.accept()

        #print("hjjhk  ", dr.switch_to.active_element.text, "  hjkhjk")
        #dr.switch_to().alert().send_keys(f"{config.jenkins_user}")
        #dr.switch_to().alert().send_keys(f"{config.jenkins_pw}")

        time.sleep(1)

        return True

class Jenkins():
    def __init__(self):
        self.tmp = 1
        """
            # job 생성
            "{config.jenkins_server}/createItem?name=[job name]"

            # job 조회
            "{config.jenkins_server}/view/{project_view}/job/{project_name}/api/json" # or xml

            # job 빌드 결과 조회
            "{config.jenkins_server}/view/{project_view}/job/{project_name}/[build number]/api/json" # or xml

            # job 빌드 결과 조회 - 마지막 성공 빌드
            "{config.jenkins_server}/view/{project_view}/job/{project_name}/lastStableBuild/api/json" # or xml

            # get job info 
            "{config.jenkins_server}/view/{project_view}/api/json"
        """

    # 입력받은 정보로 item 생성
    def create_item(self, project_view, project_name):
        pass

    # 입력받은 프로젝트 명을 가진 item의 빌드를 api로 진행
    def build(self, project_view = "test_Company", project_name=""):
        #query = f"curl -X POST {config.jenkins_server}/view/{project_view}/job/{project_name}/build --user {config.jenkins_user}:{config.jenkins_pw}"
        query = f"curl -X POST {config.jenkins_server}/job/{project_name}/build --user {config.jenkins_user}:{config.jenkins_pw}"
        
        print(query)
        os.system(query)

class Gui():
    def __init__(self):
            pass
        
    class Mouse():
        def __init__(self):
            pass

    class Keyboard():
        def __init__(self):
            pass

    def connect_targetElements(self, target_process):
        app = pywinauto.application.Application(backend="uia")
        for proc in pywinauto.findwindows.find_elements():
            if proc.name == "*" + target_process + "*":
                break

        return app.connect(process = proc.process_id)

    def ctrl_a(self):
        pyautogui.hotkey('ctrl', 'a')

    def ctrl_c(self):
        pyautogui.hotkey('ctrl', 'c')

    def ctrl_v(self):
        pyautogui.hotkey('ctrl', 'v')

    def Quick_launch(self):
        pyautogui.hotkey('win', 'r')
    
    def create_notepad(self):
        self.Quick_launch()
        pyautogui.typewrite('notepad', interval=0.1)
        pyautogui.press('enter')
        
    def get_clipboard(self):
        pass
    
    def set_focus(self):
        pass 

    def main(self):
        #my_gui = Gui()
        #my_gui.connect_targetElements("메모장").print_control_identifiers()
        #pyautogui.click(200, 100)
        #my_gui.ctrl_a()
        #my_gui.ctrl_c()

        #my_gui.create_notepad()
        #clip_text = pyperclip.paste()
        #print(clip_text)
        #print(pyautogui.size())
        print(pyautogui.position())

class Gerrit():
    def __init__(self):
        pass

    # gerrit REST API
    def send_api(self, path, method, body={}):
        host = config.gerrit_host + "/a"
        url = host + path

        headers = {'Content-Type':'application/json', 'charset':'UTF-8'}
        auth = requests.auth.HTTPDigestAuth(config.gerrit_username, config.gerrit_token)

        try:
            if method == 'GET':
                response = requests.get(url, headers=headers, auth=auth)
            elif method == 'POST':
                if body is None:
                    return
                
                response = requests.post(url, headers=headers, data=json.dumps(body, ensure_ascii=False, indent="\t"), auth=auth)
            elif method == 'PUT':
                response = requests.put(url, auth=auth)

            return json.loads(response.text[4:] if response.text.startswith(')]}\'') else response.text)

        except:
            pass

    # 사용법 및 테스트 필요
    def submit_change(self, id):
        api_query = "/changes/" + id + "/submit"
        body = {
            "wait_for_merge":True
        }
        
        try:
            json_res = self.send_api(api_query, "POST", body=body)
            
        except Exception as ex:
            print(ex)

        
    # CodeReview Check
    def Check_CodeReview(self, sec=20):
        while(sec):
            try:
                str_status, str_text = self.send_api("/changes/?q=project:projcet/rpm+reviewer:self+is:open", "GET")
                str_text = str_text.replace(")]}'", "").strip()

                #print(str_text)
                if(len(str_text) < 10):
                    print(datetime.datetime.now(), ' - 없음')
                    time.sleep(sec)

                else:
                    print(datetime.datetime.now(), ' - 있음', str_text)
                    self.msgbox("Check your code review job", "Notice : Gerrit")

                    # 코드리뷰 페이지 열어주고 소리내기
                    for i in range(5):
                        ws.beep(2000, 1000)

            except:
                pass
            
            
    # windows MessageBox  
    def msgbox(self, str_msg, str_title):
        ctypes.windll.LoadLibrary("user32.dll").MessageBoxW(0, str_msg, str_title, 1)


    def tmp_check_branchName(self, str):
        url = ""

        #return True
        return False


    def download_CsvFile(self):
        url = ""
        user = ""
        pw = ""

        with urlopen(url, auth=(user, pw)) as f:
            html = f.read().decode('utf-8')

        print(html) 
        #with open(file_name, "wb") as file:

    def get_all_siteBranch(self):
        list_hyphen_version = []
        json_tmp = self.send_api("/projects/projcet%2Frpm/branches", "GET")

        for tmp in json_tmp:
            try:
                str_tmp = tmp['ref'].split("release")[1]
                if "/" not in str_tmp or "-" not in str_tmp:
                    continue
                if "DEVOPS" in str_tmp:
                    continue
                list_hyphen_version.append("release"+str_tmp)
            except:
                pass
        
        return list_hyphen_version

    def get_all_siteCode(self):
        list_hyphen_version = []
        json_tmp = self.send_api("/projects/projcet%2Frpm/branches", "GET")

        for tmp in json_tmp:
            try:
                str_tmp = tmp['ref'].split("release")[1][1:]
                if "-" not in str_tmp:
                    continue
                if "DEVOPS" in str_tmp:
                    continue
                list_hyphen_version.append(str_tmp)
            except:
                pass
            
        dict_tmp = {}
        for x in list_hyphen_version:
            dict_tmp[x.split("-")[1]] = ""
        
        return list(dict_tmp.keys())

    #프로젝트 목록 받아오기
    def get_projects(self):
        list_project_file = []
        json_tmp = self.send_api("/projects/?d", "GET")

        for tmp in json_tmp:
            try:
                str_tmp = tmp#.split("release-")[1]
                if ("projcet/" in str_tmp) and ("UC" not in str_tmp):
                    list_project_file.append(str_tmp)
            except:
                pass

        print(list_project_file)
        return list_project_file
            
            
    # 특정 Change의 커밋 메시지를 가져옴
    def get_revisions_id(self):
        api_query = "/changes/?q=project:projcet/rpm+reviewer:self+is:open"
        list_res = []

        try:
            json_res = self.send_api(api_query, "GET")
            for json_tmp in json_res:
                list_res.append(json_tmp['id'])

        except Exception as ex:
            print(ex)

        return list_res

    def get_id(self, id):
        api_query = "/changes/" + id

        try:
            json_res = self.send_api(api_query, "GET")
            return json_res

        except Exception as ex:
            print(ex)
    
    # 특정 Change의 정보를 json으로 받아옴 (내가 올린것만)
    def get_subject(self, id):
        api_query = "/changes/" + id

        try:
            json_res = self.send_api(api_query, "GET")
            if json_res['owner']['_account_id'] == config.gerrit_user_id:
                return json_res

        except Exception as ex:
            print(ex)

    def main(self):
        gerrit = Gerrit()
    
        list_tmp = []
        for x in list_tmp:
            json_tmp = gerrit.get_id(x)
            prefix_tmp = json_tmp['project'].replace("/", "_")
            tmp = json_tmp['branch'].split("/")[1].split("-")
            suffix_tmp = tmp[0] + ".0 -> " + tmp[0] + ".1"
            site_code = tmp[1]

            print(prefix_tmp + "-" + suffix_tmp)
            #q = "/changes/" + json_tmp['id'] + "/COMMIT_MSG"
            #print(gerrit.send_api(q, "GET"))

            # 작동은 되지만 사용은 x
            #gerrit.submit_change(x)

        #gerrit.get_revisions_id()
        print("")

class Excel():
    def __init__(self):
        pass
    
    def test_sheet(self):
        sheet = pyexcel.Sheet()
        sheet.name = "name"
        sheet.ndjson = """
        {"year": ["2017", "2018", "2019", "2020", "2021"]}
        {"user": [129, 253, 304, 403, 545]}
        {"visit": [203, 403, 632, 832, 1023]}
        """.strip()

        print(sheet)

class Docx():
    def __init__(self, file_path):
        self.file_path = file_path

    def append_contents(self, paragraph):
        try:
            doc = docx.Document(self.file_path)
        except:
            doc = docx.Document()
            doc.save(self.file_path)
            doc = docx.Document(self.file_path)

        doc.add_paragraph().add_run("\n" + datetime.now().strftime('%Y-%m-%d')).bold = True
        doc.add_paragraph(paragraph)

        doc.save(self.file_path)

    def main(self):
        my_doc = Docx(path.test_desktop_path)
        for i in range(10):
            my_doc.append_contents(f"test_{i}\n")

class Test():
    import path as path_my

    def __init__(self):
        pass

    def find_dir(self, path, str_space):
        for x in os.listdir(path):
            tmp = f"{path}/{x}".replace("//", "/")
            if isdir(tmp):
                print(str_space + x)
                self.find_dir(tmp, str_space + "  ")

    def make_dir(self, path):
        for x in os.listdir(path):
            if ".ini" in x: continue
            if "기타" in x: continue

            tmp = f"{path}/{x}".replace("//", "/")
            if isdir(tmp):
                pass

    def find_path(self):
        target_path = [path_my.gdrive_path_version]

        for tmp_path in target_path:
            print("\nstart : ", tmp_path) 
            make_dir(tmp_path, "")
            print("\nend : ", tmp_path) 

    def fastapi_test(self):
        from fastapi import FastAPI
        import uvicorn
        import asyncio

        async def app(scope, receive, send):
            assert scope['type'] == 'http'

            await send({
                'type': 'http.response.start',
                'status': 200,
                'headers': [
                    [b'content-type', b'text/plain'],
                ],
            })
            await send({
                'type': 'http.response.body',
                'body': b'test site',
            })

        async def main():
            config = uvicorn.Config("test2:app", port=8080, log_level="info", reload=True)
            server = uvicorn.Server(config)
            await server.serve()

        #if __name__ == "__main__":
            #asyncio.run(main())
        def main():
            asyncio.run(main())

    # Create dummy file for TEST
    def test_file_create(data, unit="gb"):
        dict_size = {
            "b":1,
            "kb":1024, 
            "mb":1024*1024, 
            "gb":1024*1024*1024, 
            "B":1,
            "KB":1024, 
            "MB":1024*1024, 
            "GB":1024*1024*1024, 
        }
    
        with open("test_" + str(len(data)) + unit + ".txt", "wb") as f:
            f.write(data * dict_size[unit])
        
        print("Write ", len(data) * dict_size[unit])
    
        # 이스케이프 시퀀스
        # = \u003d
        # / %2F
        #   \u0020
        #test_file_create(b'1', "b") # 대용량 더미 txt 파일 생성 
        print("")


# 코드리뷰 대기 중인 change들 내용을 특정 경로에 파일로 생성해줌 
# -> 빌드 후 이것도 같이 넘기면 작성 편함
def make_detail_file():
    gerrit = Gerrit()
    date = str(datetime.datetime.now()).split(" ")[0] 
    file_path = config.path_rpm + "/" + date
    
    try:
        os.mkdir(file_path)
        
    except:
        pass
    
    for id in gerrit.get_revisions_id():
        json_tmp = gerrit.get_subject(id)
        
        if json_tmp == None:
            continue
        
        str_subject = json_tmp["subject"]
        str_branch = str(json_tmp["branch"]).replace("release/", "")
        
        # 일부 버전 확인 후 예외처리 로직 필요
        if "-" not in str_branch:
            continue
            
        try:
            file_name = file_path + "/" + str_branch + ".txt"
            file_content = str(str_subject).replace("[", "\n[")[1:]
            
            print("file : ", file_name)
            
            #with open(file_name, "w") as file:
            #    file.write(file_content)

        except:
            print(str_branch, " - err")


if __name__ == '__main__':
    jira = Jira()
    mode = 1
    if mode == 1:
        """
        나에게 할당된 이슈 오래된 순서로 정렬해서 가져오기
        - 작성일 : 2022-09-15
        - 마지막 테스트 날짜 : 2022-10-31, 성공
        - 기능 사용 대상 : 
        - 기능 사용 가능 조건 : 
        - 사용 시 주의사항 : config.py의 token과 id 를 본인의 값으로 변경해주어야함.
        """
        cnt = 0
        file_path =  path.home_path + f"/Desktop/MyIssue_{jira.today_yyyymmdd}.txt" # 파일명, 경로 변경 필요 (날짜 넣기)
        list_result = list(reversed(sorted(jira.get_my_issue(), key=lambda x:x[0])))
        
        str_res = ""
        for str_tmp in list_result:
            str_res += str(str_tmp) + "\n"
            cnt+=1
        
        str_res_2 = ""
        for str_tmp in list_result:
            str_res_2 += str(str_tmp[1]) + "\n"
            cnt+=1
        
        print("count :" + str(cnt))

        with open(file_path, "w") as f:
            f.write(str_res)
            f.write(str_res_2)

    elif mode == 2: # 이슈 속성 얻어오기
        #q = 'project = "" ORDER BY created DESC'
        q = 'key = "H0123-0000"'
        issues = jira.jira.search_issues(q)
        for iss in issues:
            #o = jira.jira.issue(iss.key)
            print(iss)
            for i in dir(iss.fields):
                print(i, str(getattr(iss.fields, i)))
                #if "" == i:
                    #print(str(getattr(iss.fields, i)))
            #    if i == 'assignee' and "담당자" in str(getattr(iss.fields,i)):
            #        print(iss.key, " - " + iss.raw['fields'][i])
            #print(o.assignee)

        #iss = jira.jira.issue('이슈코드')
        #for i in dir(iss.fields):
        #    print(i+":"+str(getattr(iss.fields,i)))
    elif mode == 3:
        print(jira.get_my_issue())
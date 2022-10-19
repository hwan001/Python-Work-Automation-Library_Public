import sys
import subprocess

try:
    import docx

except:
    subprocess.check_call([sys.executable,'-m', 'pip', 'install', '--upgrade', 'pip'])
    subprocess.check_call([sys.executable,'-m', 'pip', 'install', '-r', 'requirements.txt'])

    import docx
from datetime import datetime

import path

class Docx():
    def __init__(self, file_path):
        self.file_path = file_path

    def append_contents(self, paragraph):
        try:
            doc = docx.Document(self.file_path)
        except:
            doc = docx.Document()
            doc.save(self.file_path)
            doc = docx.Document(self.file_path)

        doc.add_paragraph().add_run("\n" + datetime.now().strftime('%Y-%m-%d')).bold = True
        doc.add_paragraph(paragraph)

        doc.save(self.file_path)

if __name__ == '__main__':
    my_doc = Docx(path.test_desktop_path)
    for i in range(10):
        my_doc.append_contents(f"test_{i}\n")